"""Generate the project paper as a DOCX using the provided Word template."""

from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.image as mpimg
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


PROJECT_ROOT = Path(__file__).resolve().parents[3]
EVALUATION_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = Path("/Users/saroshnadaf/Downloads/Paper_Format.docx")
OUTPUT_PATH = EVALUATION_DIR / "fault_tolerant_financial_transcript_analysis.docx"
ASSETS_DIR = EVALUATION_DIR / "paper_docx_assets"
EVAL_FIGURES_DIR = EVALUATION_DIR / "finbert_intent_v3_eval"


TITLE = "Fault-Tolerant Financial Transcript Analysis with Lazy Model Orchestration Under CPU-Constrained Cloud Deployment"
AUTHORS = "Author One, Author Two, Author Three"
AFFILIATIONS = [
    "Affiliation Placeholder",
    "Email Placeholder",
]
KEYWORDS = "Financial NLP, FinBERT, fault-tolerant inference, FastAPI, lazy loading, Railway deployment"


ABSTRACT = (
    "This paper presents a production-oriented financial transcript analysis system designed for constrained cloud "
    "deployment rather than benchmark-only model evaluation. The system processes earnings-call transcripts through "
    "a FastAPI backend that performs role-aware segmentation, segment-level FinBERT intent inference, "
    "conversation-level aggregation, and market-oriented prediction. Its principal contribution is a fault-tolerant "
    "inference architecture that preserves API availability under memory pressure, dependency instability, and "
    "fallback-model failures in a CPU-only Railway environment. The backend uses Hugging Face-hosted models for the "
    "primary intent classifier and a secondary PyTorch fallback model that is loaded lazily, guarded by a thread "
    "lock, protected by a timeout, and disabled permanently by a circuit breaker after failure. This design "
    "eliminates startup crashes and avoids repeated memory spikes associated with eager deserialization of large .pt "
    "artifacts. To stabilize deployment, the runtime is pinned to Python 3.11.9 with CPU-only PyTorch and "
    "compatible versions of Transformers, NumPy, and auxiliary libraries. Internal evaluation on a balanced sample "
    "of 240 transcripts shows that the proposed system achieves an accuracy of 0.8292 and a macro F1 score of "
    "0.8303, compared with 0.5042 and 0.4524 for a FinBERT baseline. The results indicate that deployment-aware ML "
    "system design can yield both stronger practical robustness and improved transcript-level reasoning."
)


SECTIONS = [
    (
        "I. Introduction",
        [
            "Financial transcript analysis is a demanding natural language processing task because it must convert "
            "long, speaker-rich earnings-call transcripts into structured insights suitable for downstream decision "
            "support. In this project, the system analyzes transcript text through an end-to-end pipeline that "
            "combines segmentation, intent classification, score aggregation, signal generation, volatility "
            "estimation, and market prediction. Unlike purely model-centric work, the central problem addressed here "
            "is operational reliability: the service must remain available in a CPU-only, memory-limited cloud "
            "environment where dependency conflicts and model-loading overhead can easily turn a functional model "
            "into an unstable production system.",
            "The deployed service is built around a FastAPI backend exposing /analyze, /upload, and /compare. The "
            "primary inference path uses a Hugging Face-hosted FinBERT-based intent model, while a secondary PyTorch "
            "fallback model is retained to preserve functionality when the primary path or optional components fail. "
            "The system operates under Railway deployment constraints, where container size and startup stability "
            "matter as much as classification accuracy. To satisfy those constraints, the codebase uses CPU-only "
            "runtime pinning, lazy model loading, and explicit circuit-breaking behavior when fallback "
            "initialization fails.",
            "The contributions of this work are threefold. First, it introduces a multi-stage financial NLP pipeline "
            "that transforms transcript segments into intent labels, aggregate scores, and directional market "
            "signals. Second, it implements a fault-tolerant model-serving layer with thread-safe lazy "
            "initialization, timeout-guarded loading, and safe degradation so that API requests never crash even "
            "when fallback inference fails. Third, it demonstrates a deployment-aware engineering strategy for "
            "CPU-only cloud execution by moving model artifacts to Hugging Face and pinning the runtime dependency "
            "set to prevent the class of incompatibilities that previously caused startup failures and image-size "
            "blowups.",
        ],
    ),
    (
        "II. Related Work",
        [
            "Transformer models established the foundation for modern language understanding by replacing recurrent "
            "sequence processing with self-attention mechanisms [3]. BERT further advanced this paradigm through "
            "bidirectional pretraining, enabling strong downstream performance across a range of NLP tasks [1]. In "
            "financial NLP, FinBERT adapted BERT-style representations to finance-specific language and showed that "
            "domain pretraining improves sentiment and contextual interpretation in financial documents [2].",
            "The present system differs from prior model-centric literature by focusing on production orchestration. "
            "Hugging Face provides the model distribution layer that enables pretrained checkpoints to be hosted and "
            "loaded programmatically across environments [4], [17]. PyTorch provides the serialization and runtime "
            "machinery required to serve custom fallback models, but its flexibility introduces memory and "
            "compatibility risks when large serialized objects are loaded on constrained hardware [5]. FastAPI "
            "supplies the lightweight asynchronous web layer that makes these models accessible as a service [6]. "
            "This work combines these technologies into a fault-tolerant serving pipeline that remains operational "
            "under failure conditions, which is a practical concern not addressed by many benchmark-driven studies.",
            "Broader work on pretrained language models further contextualizes the design space explored by this "
            "system. Large-scale few-shot models expanded the general-purpose capabilities of transformer-based NLP "
            "[7], while robust BERT variants such as RoBERTa, XLNet, and ELECTRA improved pretraining efficiency "
            "and representation quality [8], [9], [10]. Domain-adapted models such as BioBERT and LEGAL-BERT "
            "demonstrated that specialization to a professional corpus can significantly improve downstream "
            "performance [11], [13]. Long-document and sequence-to-sequence architectures such as Longformer and "
            "BART further extended transformer utility for extended-context processing and generative settings [12], "
            "[14]. Related representation learning approaches including Sentence-BERT and ULMFiT also informed "
            "practical strategies for downstream adaptation and semantic modeling [15], [16].",
        ],
    ),
    (
        "III. System Architecture",
        [
            "The system architecture centers on a FastAPI backend that accepts transcript text or uploaded files and "
            "returns a structured analysis payload. The active backend entrypoint is backend/api/server.py, which "
            "exposes the /analyze, /upload, and /compare routes. The backend is stateless in its current design, "
            "and CORS is enabled for development use. The primary data flow begins with transcript ingestion, "
            "continues through segmentation and intent inference, and ends with aggregation into score, signal, "
            "confidence, volatility, prediction, and driver fields.",
            "The first model layer is the TranscriptAnalyzer, which produces a list of transcript segments. It uses "
            "speaker-aware parsing when cues such as CEO:, CFO:, ANALYST:, or OPERATOR: are present and falls back "
            "to sentence chunking when structured speaker markers are absent. Each segment is sent to the intent "
            "model, which is a Hugging Face-hosted four-class FinBERT sequence classifier. The intent output is then "
            "consumed by the signal engine, which maps segment intents to numeric scores and derives a "
            "transcript-level signal. A market predictor converts that signal and its support statistics into a "
            "directional forecast and explanation. A separate insight layer extracts growth and risk drivers from "
            "the transcript text. The response returned by /analyze includes score, signal, prediction, confidence, "
            "volatility, intent_distribution, growth_drivers, risk_drivers, timeline, and fallback_used.",
            "A critical architectural feature is the fallback model orchestration. If the FinBERT path fails, the "
            "analyzer can load a secondary PyTorch transformer from a Hugging Face-hosted .pt artifact. That "
            "fallback is not loaded at startup. Instead, it is loaded lazily under a thread-safe lock, with a "
            "strict timeout and a process-level circuit breaker. This design ensures that model failures remain "
            "localized and do not crash the entire API process.",
        ],
    ),
    (
        "IV. Methodology",
        [],
    ),
    (
        "V. Deployment and Optimization",
        [
            "The system is deployed on Railway using a CPU-only configuration. This deployment choice shaped the "
            "entire runtime strategy. A GPU-enabled PyTorch package would have increased image size substantially "
            "and risked exceeding deployment limits, so the runtime is pinned to a CPU-only wheel. The repository "
            "also includes a Python runtime pin of 3.11.9 to reduce interpreter drift.",
            "The current backend dependency set is intentionally minimal and stable. PyTorch is pinned to a CPU-only "
            "build, Transformers is pinned to a version compatible with that Torch release, NumPy is constrained "
            "below version 2 to avoid ABI issues, and auxiliary packages such as Pandas, Datasets, "
            "python-multipart, pdfplumber, and protobuf are pinned to versions observed to be compatible with the "
            "current runtime. This pinning strategy prevents resolver drift and the cascading import failures "
            "previously observed in deployment.",
            "The model artifacts themselves were migrated away from repository storage and are now hosted on Hugging "
            "Face. This decision was necessary because local GitHub storage and bandwidth constraints made it "
            "impractical to keep large model binaries in the repository. As a result, runtime code downloads "
            "checkpoints on demand rather than shipping them inside the application image. The primary intent model "
            "is loaded from a Hugging Face repository, and the fallback .pt artifact is also fetched from Hugging "
            "Face when required. This separation keeps the deployment image small while preserving functionality.",
            "The most important deployment optimization is the removal of eager fallback loading. A fallback .pt "
            "file can temporarily double memory usage during deserialization and trigger Railway container restarts "
            "if loaded at startup. By moving the fallback to lazy loading with timeout and circuit breaker "
            "protection, the system avoids startup crashes and keeps /analyze available even under resource "
            "pressure. In practice, this is the difference between a fragile prototype and a service that can "
            "survive production traffic.",
        ],
    ),
    (
        "VI. Results and Evaluation",
        [
            "Evaluation is implemented in backend/financial_pragmatic_ai/evaluation/better_than_fin/evaluate.py and "
            "runs on a balanced sample built from pragmatic_intent_dataset_clean.csv. The latest saved benchmark "
            "uses 240 transcripts total, with 80 examples per signal class. The comparison is between a FinBERT "
            "sentiment baseline and the proposed multi-stage system.",
            "The proposed system improves accuracy by 32.50 percentage points and macro F1 by 37.79 percentage "
            "points relative to the FinBERT baseline. This improvement is not solely attributable to the classifier, "
            "but to the system-level reasoning pipeline that maps segment intents into financial signals and then "
            "into market predictions. The baseline directly predicts polarity-like output, while the proposed system "
            "uses transcript structure, speaker-aware segmentation, and intent aggregation to extract more aligned "
            "financial meaning.",
            "The prediction distribution also indicates better class utilization. In the latest evaluation, the "
            "FinBERT baseline produced neutral: 120, growth: 102, and risk: 18, whereas the proposed system "
            "produced growth: 61, risk: 67, and neutral: 112. The custom system therefore shows stronger "
            "sensitivity to risk language than the baseline and avoids the extreme collapse toward a single class "
            "that was observed earlier in development. This behavior is consistent with the explicit "
            "intent-to-signal mapping and aggregation thresholds currently used in the codebase.",
            "From a deployment standpoint, the system also behaves more robustly under failure. Fallback loading "
            "failures do not crash the API, and the response path includes a fallback_used indicator for "
            "observability. Startup succeeds without loading all model artifacts into memory, and the service "
            "remains responsive under CPU-only constraints. These properties are not captured by accuracy alone but "
            "are essential to the real-world value of the system.",
        ],
    ),
    (
        "VII. Discussion",
        [
            "The principal result of this project is that production reliability and ML quality are not separate "
            "concerns. The evaluation metrics show that the system can outperform a FinBERT baseline on the internal "
            "benchmark, but the more important engineering achievement is that the service remains available when "
            "model artifacts fail to load, when optional components are absent, and when container memory is "
            "limited. The fallback architecture is therefore a first-class contribution rather than a defensive "
            "afterthought.",
            "The system also illustrates the tradeoff between model flexibility and operational safety. PyTorch "
            "serialization is powerful, but it can be dangerous in memory-constrained environments because loading a "
            "model can temporarily require substantially more memory than the final in-memory footprint. The "
            "thread-safe lazy loader and circuit breaker address this directly. The design does not attempt to "
            "recover endlessly from failure; it disables unstable behavior and keeps the request path healthy. That "
            "is the correct choice for a public-facing inference service.",
            "There remain limitations. The current evaluation is an internal benchmark and should not be interpreted "
            "as final generalization performance because the dataset lineage and the train/eval split strategy "
            "warrant further scrutiny. The current codebase also retains some verbose debug logging and still "
            "contains a few threshold inconsistencies in higher-level insight logic. In addition, the optional "
            "conversation-attention path is present in the architecture but is not always backed by a local weight "
            "file, which means the system can still degrade to rule-based aggregation in some deployments. Future "
            "work should focus on external validation, stricter data separation, and a more unified thresholding "
            "policy across signal and insight modules.",
        ],
    ),
    (
        "VIII. Conclusion",
        [
            "This paper presented a practical financial NLP system for earnings-call analysis that emphasizes "
            "reliability, observability, and deployment stability in addition to predictive accuracy. The system "
            "combines FastAPI serving, Hugging Face-hosted FinBERT inference, lazy fallback orchestration, and a "
            "multi-stage transcript reasoning pipeline to produce structured market signals from raw transcript "
            "text. Its core contribution is a fault-tolerant serving architecture that uses CPU-only deployment, "
            "thread-safe lazy loading, timeout protection, and a circuit breaker to ensure that the API never "
            "crashes even when fallback model loading fails. Internal evaluation shows that the proposed system "
            "significantly outperforms a FinBERT baseline on accuracy and macro F1, while also providing more "
            "balanced signal behavior and better practical robustness. The result is a deployment-aware ML system "
            "that is appropriate for resource-constrained cloud environments and that demonstrates how engineering "
            "rigor can materially improve the usefulness of financial NLP models.",
        ],
    ),
]


SUBSECTIONS = [
    (
        "A. Transcript Segmentation",
        "The input transcript is first normalized and segmented so that the model does not treat an entire earnings "
        "call as a single atomic document. Speaker-aware parsing is preferred because earnings calls are structured "
        "around management and analyst turns, and these turns often carry distinct financial semantics. When "
        "explicit speaker labels are absent, the system falls back to sentence-based chunking and dynamic "
        "regrouping so that transcripts still produce multiple segments. This segmentation stage is essential "
        "because the downstream scoring logic depends on variation across segment-level intents.",
    ),
    (
        "B. Intent Prediction",
        "Each segment is classified into one of four intents: EXPANSION, COST_PRESSURE, STRATEGIC_PROBING, or "
        "GENERAL_UPDATE. The primary classifier is a FinBERT-based sequence classifier loaded from a Hugging Face "
        "repository. The model returns logits, confidence, and a CLS embedding. The intent label is then mapped "
        "into a financial interpretation that the aggregation layer can consume. This design is intentionally "
        "pragmatic: it uses domain-specialized language understanding but keeps the output space small enough to "
        "support robust aggregation and explanation.",
    ),
    (
        "C. Score Aggregation",
        "Let y_i denote the intent assigned to segment i, and let w(.) denote the intent weight mapping. The "
        "current mapping is w(EXPANSION) = +1.0, w(COST_PRESSURE) = -1.0, w(STRATEGIC_PROBING) = +0.2, and "
        "w(GENERAL_UPDATE) = 0.0. For a transcript with N segments, the aggregate score is computed as s = "
        "(1/N) sum from i=1 to N of w(y_i). The final signal is obtained by thresholding the aggregate score as "
        "growth if s > 0.2, risk if s < -0.2, and neutral otherwise. Confidence is derived from the dominant "
        "signal share, while volatility is estimated from the dispersion of mapped segment values. This rule-based "
        "aggregation provides transparency and makes the transcript-level decision easier to interpret than a "
        "monolithic classifier would.",
    ),
    (
        "D. Market Prediction and Driver Extraction",
        "The final signal, score, volatility, and intent distribution are passed to a market predictor that "
        "generates a directional outcome and a short explanation. In parallel, the insight layer extracts "
        "representative growth and risk drivers by compressing candidate sentences, deduplicating them "
        "semantically, and retaining only high-quality spans. This produces a compact driver set that can be "
        "surfaced directly in the API response and frontend UI.",
    ),
    (
        "E. Fault-Tolerant Fallback Orchestration",
        "The fallback PyTorch model is loaded only when needed. The loader executes on CPU, uses best-effort "
        "memory-friendly deserialization options, and is guarded by a lock so that only one request can attempt "
        "initialization at a time. A timeout prevents request threads from waiting indefinitely. If loading fails "
        "or times out, the circuit breaker permanently disables further attempts for the current process. This "
        "policy intentionally favors service continuity over fallback completeness. The API therefore continues "
        "returning responses even when the fallback path is unavailable.",
    ),
]


REFERENCES = [
    '[1] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding," in Proc. NAACL-HLT, 2019, pp. 4171-4186.',
    '[2] D. Araci, "FinBERT: Financial Sentiment Analysis with Pre-trained Language Models," arXiv:1908.10063, 2019.',
    '[3] A. Vaswani et al., "Attention Is All You Need," in Advances in Neural Information Processing Systems, vol. 30, 2017.',
    '[4] Hugging Face, "Transformers Documentation," 2026. [Online]. Available: https://huggingface.co/docs/transformers',
    '[5] PyTorch, "PyTorch Documentation," 2026. [Online]. Available: https://pytorch.org/docs/stable/index.html',
    '[6] FastAPI, "FastAPI Documentation," 2026. [Online]. Available: https://fastapi.tiangolo.com/',
    '[7] T. Brown et al., "Language Models are Few-Shot Learners," NeurIPS, 2020.',
    '[8] Y. Liu et al., "RoBERTa: A Robustly Optimized BERT Pretraining Approach," arXiv:1907.11692, 2019.',
    '[9] Z. Yang et al., "XLNet: Generalized Autoregressive Pretraining," NeurIPS, 2019.',
    '[10] K. Clark et al., "ELECTRA: Pre-training Text Encoders as Discriminators," ICLR, 2020.',
    '[11] J. Lee et al., "BioBERT: a pre-trained biomedical language model," Bioinformatics, 2020.',
    '[12] A. Beltagy et al., "Longformer: The Long-Document Transformer," arXiv:2004.05150, 2020.',
    '[13] I. Chalkidis et al., "LEGAL-BERT: The Muppets straight out of Law School," EMNLP, 2020.',
    '[14] M. Lewis et al., "BART: Denoising Sequence-to-Sequence Pre-training," ACL, 2020.',
    '[15] N. Reimers and I. Gurevych, "Sentence-BERT," EMNLP, 2019.',
    '[16] J. Howard and S. Ruder, "ULMFiT," ACL, 2018.',
    '[17] Hugging Face, "Hugging Face Hub Documentation," 2026.',
    '[18] Docker, "Docker Documentation," 2026.',
    '[19] Railway, "Railway Deployment Docs," 2026.',
    '[20] NumPy, "NumPy Documentation," 2026.',
    '[21] Pandas, "Pandas Documentation," 2026.',
    '[22] Pydantic, "Pydantic Documentation," 2026.',
    '[23] Starlette, "Starlette ASGI Framework," 2026.',
    '[24] Uvicorn, "ASGI Server Documentation," 2026.',
    '[25] Supabase, "Supabase Documentation," 2026.',
    '[26] OpenAI, "GPT Models Overview," 2026.',
]


def _clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _add_body_paragraph(doc: Document, text: str, style: str = "IEEE Paragraph") -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)


def _add_centered_paragraph(doc: Document, text: str, style: str = "IEEE Paragraph") -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)


def _add_abstract(doc: Document) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    label = paragraph.add_run("Abstract - ")
    label.bold = True
    paragraph.add_run(ABSTRACT)

    doc.add_paragraph("", style="Normal")

    paragraph = doc.add_paragraph(style="Normal")
    label = paragraph.add_run("Index Terms - ")
    label.bold = True
    paragraph.add_run(KEYWORDS)


def _create_placeholder_image(path: Path, title: str, detail: str) -> None:
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.axis("off")
    ax.add_patch(plt.Rectangle((0.05, 0.08), 0.9, 0.84, fill=False, linewidth=2, transform=ax.transAxes))
    ax.text(0.5, 0.62, title, ha="center", va="center", fontsize=18, weight="bold", transform=ax.transAxes)
    ax.text(0.5, 0.40, detail, ha="center", va="center", fontsize=11, wrap=True, transform=ax.transAxes)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def _create_evaluation_panel(output_path: Path) -> None:
    image_specs = [
        ("Model Comparison", EVAL_FIGURES_DIR / "model_comparison.png"),
        ("Per-Class F1", EVAL_FIGURES_DIR / "per_class_f1.png"),
        ("Agreement", EVAL_FIGURES_DIR / "agreement_bar.png"),
        ("FinBERT Confusion Matrix", EVAL_FIGURES_DIR / "confusion_matrix_finbert_normalized.png"),
        ("Our System Confusion Matrix", EVAL_FIGURES_DIR / "confusion_matrix_ours_normalized.png"),
        ("Class Distribution", EVAL_FIGURES_DIR / "class_distribution.png"),
    ]

    fig, axes = plt.subplots(2, 3, figsize=(15, 9))
    for ax, (title, path) in zip(axes.flat, image_specs):
        image = mpimg.imread(path)
        ax.imshow(image)
        ax.set_title(title, fontsize=11)
        ax.axis("off")

    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def _add_figure(doc: Document, image_path: Path, caption: str, width: Inches) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width)

    caption_paragraph = doc.add_paragraph(style="Normal")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.add_run(caption)


def _add_results_table(doc: Document) -> None:
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ["Model", "Accuracy", "Macro F1"],
        ["FinBERT baseline", "0.5042", "0.4524"],
        ["Proposed system", "0.8292", "0.8303"],
    ]

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    caption = doc.add_paragraph(style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run("Table 1. Latest Internal Evaluation Results on Balanced Sample (N = 240)")


def build_docx() -> Path:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    system_architecture_img = ASSETS_DIR / "system_architecture_placeholder.png"
    fallback_flow_img = ASSETS_DIR / "fallback_flow_placeholder.png"
    evaluation_panel_img = ASSETS_DIR / "evaluation_metrics_panel.png"

    _create_placeholder_image(
        system_architecture_img,
        "System Architecture Diagram",
        "Transcript ingestion -> segmentation -> intent prediction -> aggregation -> prediction -> API response",
    )
    _create_placeholder_image(
        fallback_flow_img,
        "Fallback Flow Diagram",
        "Primary FinBERT path -> lazy fallback load -> timeout/circuit breaker -> graceful degradation",
    )
    _create_evaluation_panel(evaluation_panel_img)

    doc = Document(str(TEMPLATE_PATH))
    _clear_document(doc)

    title = doc.add_paragraph(style="paper title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)

    authors = doc.add_paragraph(style="Author")
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run(AUTHORS)

    for line in AFFILIATIONS:
        paragraph = doc.add_paragraph(style="Affiliation")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line)

    separator = doc.add_paragraph(style="Affiliation")
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.add_run("_" * 110)

    _add_abstract(doc)

    separator = doc.add_paragraph(style="Normal")
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.add_run("_" * 110)
    doc.add_paragraph("", style="Normal")

    for heading, paragraphs in SECTIONS:
        doc.add_paragraph(heading, style="Heading 1")

        if heading == "IV. Methodology":
            for subheading, body in SUBSECTIONS:
                doc.add_paragraph(subheading, style="IEEE Heading 2")
                _add_body_paragraph(doc, body)
                if subheading == "E. Fault-Tolerant Fallback Orchestration":
                    _add_figure(
                        doc,
                        fallback_flow_img,
                        "Figure 2. Fallback orchestration flow showing primary inference, lazy fallback loading, timeout handling, circuit breaker activation, and graceful degradation.",
                        Inches(3.1),
                    )
            continue

        for paragraph in paragraphs:
            _add_body_paragraph(doc, paragraph)

        if heading == "III. System Architecture":
            _add_figure(
                doc,
                system_architecture_img,
                "Figure 1. End-to-end system architecture showing transcript ingestion, segmentation, intent prediction, aggregation, and API response generation.",
                Inches(3.1),
            )

        if heading == "VI. Results and Evaluation":
            _add_results_table(doc)
            _add_figure(
                doc,
                evaluation_panel_img,
                "Figure 3. Evaluation metric diagrams for the FinBERT baseline and the proposed system.",
                Inches(3.2),
            )

    doc.add_paragraph("References", style="Heading 1")
    for reference in REFERENCES:
        doc.add_paragraph(reference, style="Normal")

    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_docx()
    print(output)
